# ============================================
# SISTEMA COMPLETO: BITÁCORA POR VOZ AUTOMÁTICA
# ============================================
import os
import json
import uuid
import whisper
import datetime
import pytz
from docx import Document
from apscheduler.schedulers.background import BackgroundScheduler
from telegram.ext import Updater, MessageHandler, CommandHandler, Filters

# ================= CONFIG =================
TELEGRAM_TOKEN = os.environ["8668132168:AAFMfXijRqQ3fWHMI9zvO2BKuNZqZRqtShQ"]
CHAT_ID = int(os.environ["572066566"])
ZONA = pytz.timezone("America/Matamoros")
BITACORA_FILE = "bitacora_data.json"

model = whisper.load_model("base")

# ================= PERSISTENCIA =================
def cargar_bitacora():
    if os.path.exists(BITACORA_FILE):
        with open(BITACORA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def guardar_bitacora_disco(bitacora):
    with open(BITACORA_FILE, "w", encoding="utf-8") as f:
        json.dump(bitacora, f, ensure_ascii=False, indent=2)

bitacora = cargar_bitacora()

# ================= FUNCIONES =================
def hoy_fecha():
    return datetime.datetime.now(ZONA).strftime("%Y-%m-%d")

def es_dia_laboral():
    return datetime.datetime.now(ZONA).weekday() < 5

def periodo_actual():
    dia = datetime.datetime.now(ZONA).day
    return "12-25" if 12 <= dia <= 25 else "26-11"

def transcribir_audio(file_path):
    result = model.transcribe(file_path)
    return result["text"]

def procesar_texto(texto):
    import re
    partes = re.split(r"\s+y\s+|,\s*", texto)
    return [
        {
            "descripcion": p.strip().capitalize(),
            "resultado": "Actividad realizada correctamente"
        }
        for p in partes if p.strip()
    ]

def registrar_actividades(fecha, actividades):
    if fecha not in bitacora:
        bitacora[fecha] = {"horas": "2:00", "actividades": []}
    bitacora[fecha]["actividades"].extend(actividades)
    guardar_bitacora_disco(bitacora)

def generar_word():
    doc = Document()
    doc.add_heading('REPORTE DE ACTIVIDADES ADMINISTRATIVAS', 0)
    doc.add_paragraph('NOMBRE: ING. RAUL DANIEL GONZALEZ IRACHETA')
    doc.add_paragraph(f'PERIODO: {periodo_actual()}')

    for fecha, datos in sorted(bitacora.items()):
        doc.add_paragraph(f"Fecha: {fecha}")
        doc.add_paragraph(f"Horas: {datos['horas']}")
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        for i, h in enumerate(["No.", "Descripción", "Resultados"]):
            table.rows[0].cells[i].text = h
        for i, act in enumerate(datos["actividades"], 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = act["descripcion"]
            row[2].text = act["resultado"]
        doc.add_paragraph("")

    filename = f"bitacora_{periodo_actual()}.docx"
    doc.save(filename)
    return filename

# ================= HANDLERS TELEGRAM =================
def recibir_audio(update, context):
    file_path = f"audio_{uuid.uuid4().hex}.ogg"
    try:
        file = update.message.voice.get_file()
        file.download(file_path)

        texto = transcribir_audio(file_path)
        actividades = procesar_texto(texto)
        fecha = hoy_fecha()
        registrar_actividades(fecha, actividades)

        resumen = "\n".join(f"• {a['descripcion']}" for a in actividades)
        update.message.reply_text(f"✅ Registrado para {fecha}:\n{resumen}")

    except Exception as e:
        update.message.reply_text(f"❌ Error al procesar el audio: {str(e)}")
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

def cmd_generar(update, context):
    if not bitacora:
        update.message.reply_text("📭 No hay actividades registradas aún.")
        return
    try:
        filename = generar_word()
        with open(filename, "rb") as f:
            context.bot.send_document(
                chat_id=update.effective_chat.id,
                document=f,
                filename=filename,
                caption=f"📄 Bitácora periodo {periodo_actual()}"
            )
    except Exception as e:
        update.message.reply_text(f"❌ Error al generar el reporte: {str(e)}")

def cmd_resumen(update, context):
    fecha = hoy_fecha()
    if fecha not in bitacora or not bitacora[fecha]["actividades"]:
        update.message.reply_text(f"📭 Sin registros para hoy ({fecha}).")
        return
    acts = bitacora[fecha]["actividades"]
    resumen = "\n".join(f"{i}. {a['descripcion']}" for i, a in enumerate(acts, 1))
    update.message.reply_text(f"📋 Actividades del {fecha}:\n{resumen}")

def cmd_start(update, context):
    update.message.reply_text(
        "👋 ¡Hola! Soy tu bot de bitácora.\n\n"
        "📌 Comandos disponibles:\n"
        "• Envía una 🎤 *nota de voz* para registrar actividades\n"
        "• /resumen — Ver actividades de hoy\n"
        "• /generar — Descargar el reporte Word\n\n"
        "⏰ Te recordaré a las 4:30 PM si no has registrado nada.",
        parse_mode="Markdown"
    )

# ================= RECORDATORIO =================
def verificar_bitacora(bot):
    if not es_dia_laboral():
        return
    fecha = hoy_fecha()
    if fecha not in bitacora or not bitacora[fecha]["actividades"]:
        bot.send_message(
            chat_id=CHAT_ID,
            text="⚠️ No has registrado bitácora hoy. Envía una nota de voz."
        )

# ================= MAIN =================
updater = Updater(TELEGRAM_TOKEN, use_context=True)
dp = updater.dispatcher

dp.add_handler(CommandHandler("start", cmd_start))
dp.add_handler(CommandHandler("generar", cmd_generar))
dp.add_handler(CommandHandler("resumen", cmd_resumen))
dp.add_handler(MessageHandler(Filters.voice, recibir_audio))

scheduler = BackgroundScheduler()
scheduler.add_job(
    verificar_bitacora,
    'cron',
    hour=16, minute=30,
    timezone=ZONA,
    args=[updater.bot]
)
scheduler.start()

print("✅ Bot iniciado. Esperando mensajes...")
updater.start_polling()
updater.idle()
